# Import basic libraries
from llama_index.readers.file import PandasExcelReader
from llama_index.core import Document
from docx import Document as DocxDocument
from paddleocr import PaddleOCR
import fitz
import os
import pandas as pd


def read_documents(path: str, ocr_treshold=0) -> list:
    ocr = PaddleOCR(use_angle_cls=True, lang='id', show_log=False)
    reader = PandasExcelReader()

    docs_all = []

    document_path = path
    df_metadata = pd.read_csv('./files_metadata.csv')

    for file in os.listdir(document_path):
        file_path = os.path.join(document_path, file)
        file_metadata = df_metadata[df_metadata['new_filename'] == file]
        if file.endswith('.pdf'):
            text = extract_text_from_pdf(file_path, ocr, treshold=ocr_treshold)
            document = Document(
                text=text,
                metadata={
                    "file_name": file_metadata['new_filename'].values[0],
                    "title": file_metadata['title'].values[0],
                    "sector": file_metadata['sektor'].values[0],
                    "subsector": file_metadata['subsektor'].values[0],
                    "regulation_type": file_metadata['jenis_regulasi'].values[0],
                    "regulation_number": file_metadata['nomor_regulasi'].values[0],
                    "effective_date": file_metadata['tanggal_berlaku'].values[0],
                }
            )
            docs_all.append(document)
        elif file.endswith('.xlsm') or file.endswith('.xlsx') or file.endswith('.xls'):
            # Handle for Excel (xlsm, xlsx)
            text = extract_from_excel(file_path, reader=reader)
            document = Document(
                text=text,
                metadata={
                    "file_name": file_metadata['new_filename'].values[0],
                    "title": file_metadata['title'].values[0],
                    "sector": file_metadata['sektor'].values[0],
                    "subsector": file_metadata['subsektor'].values[0],
                    "regulation_type": file_metadata['jenis_regulasi'].values[0],
                    "regulation_number": file
                    ['nomor_regulasi'].values[0],
                    "effective_date": file_metadata['tanggal_berlaku'].values[0],
                }
            )

            text = text.replace('\n', ' ')
            docs_all.append(document)

        elif file.endswith('.docx'):
            # Handle for Word (docx)
            text = extract_from_docx(file_path)
            document = Document(
                text=text,
                metadata={
                    "file_name": file_metadata['new_filename'].values[0],
                    "title": file_metadata['title'].values[0],
                    "sector": file_metadata['sektor'].values[0],
                    "subsector": file_metadata['subsektor'].values[0],
                    "regulation_type": file_metadata['jenis_regulasi'].values[0],
                    "regulation_number": file_metadata['nomor_regulasi'].values[0],
                    "effective_date": file_metadata['tanggal_berlaku'].values[0],
                }
            )

            docs_all.append(document)
    print(f"Read {len(docs_all)} documents")

    return docs_all


# =============================================================================

def extract_text_and_images_from_page(doc, page, ocr, treshold):
    text = page.get_text()
    image_text = ""
    image_list = page.get_images(full=True)
    # Iterate through all images found on the page
    for image_info in image_list:
        xref = image_info[0]
        image_dict = doc.extract_image(xref)
        image_bytes = image_dict['image']
        # Use PaddleOCR to extract text from the image
        ocr_result = ocr.ocr(image_bytes)
        # Check if OCR result is valid before processing
        if ocr_result and ocr_result != [None]:
            for result in ocr_result:
                for res in result:
                    text_tuple = res[1]
                    text_string = text_tuple[0]
                    text_confidence = text_tuple[1]  # For confidence threshold
                    if text_confidence > treshold:
                        image_text += text_string + '\n'
    # Combine page text and image text
    return text + "\n" + image_text

def extract_text_from_pdf(file_path, ocr, treshold):
    # Load the PDF file
    doc = fitz.open(file_path)
    text = ""
    # Iterate through all pages in the PDF
    for page_num in range(len(doc)):
        page = doc[page_num]
        # Extract text and images from the page
        page_text = extract_text_and_images_from_page(doc, page, ocr, treshold)
        text += page_text + "\n"
    return text

# =============================================================================

def extract_from_docx(file_path):
    doc = DocxDocument(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

# =============================================================================

def extract_from_excel(file_path, reader=PandasExcelReader()):
    doc = reader.load_data(file_path)
    text = ''
    for sheet in doc:
        for row in sheet:
            for cell in row:
                text += str(cell) + ' '
    return text
